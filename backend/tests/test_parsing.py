import pytest

from app.services.parsing import UnsupportedFileType, parse_document


def test_parse_txt():
    pages = parse_document("politique.txt", "Politique d'usage de l'IA — accès contrôlé.".encode())
    assert pages == ["Politique d'usage de l'IA — accès contrôlé."]


def test_parse_md():
    pages = parse_document("charte.md", b"# Charte\nGouvernance de l'IA")
    assert len(pages) == 1
    assert "Gouvernance" in pages[0]


def test_parse_pdf_per_page():
    import fitz

    doc = fitz.open()
    for text in ["Page une : gouvernance des données.", "Page deux : cycle de vie du modèle."]:
        page = doc.new_page()
        page.insert_text((72, 72), text)
    data = doc.tobytes()

    pages = parse_document("politique.pdf", data)
    assert len(pages) == 2
    assert "gouvernance" in pages[0]
    assert "cycle de vie" in pages[1]


def test_parse_docx():
    import io

    from docx import Document as DocxDocument

    d = DocxDocument()
    d.add_paragraph("Politique de gestion des risques IA.")
    buf = io.BytesIO()
    d.save(buf)

    pages = parse_document("risques.docx", buf.getvalue())
    assert len(pages) == 1
    assert "risques IA" in pages[0]


def test_unsupported_type():
    with pytest.raises(UnsupportedFileType):
        parse_document("image.png", b"\x89PNG")


def test_parse_docx_preserves_table_position():
    import io

    from docx import Document as DocxDocument

    d = DocxDocument()
    d.add_paragraph("AVANT le tableau.")
    table = d.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "MILIEU-A"
    table.rows[0].cells[1].text = "MILIEU-B"
    d.add_paragraph("APRES le tableau.")
    buf = io.BytesIO()
    d.save(buf)

    text = parse_document("ordre.docx", buf.getvalue())[0]
    assert text.index("AVANT") < text.index("MILIEU-A") < text.index("APRES")


def test_empty_txt_rejected():
    from app.services.parsing import EmptyDocument

    with pytest.raises(EmptyDocument):
        parse_document("vide.txt", b"   \n\t  ")


def test_blank_pdf_rejected():
    import fitz

    from app.services.parsing import EmptyDocument

    doc = fitz.open()
    doc.new_page()
    with pytest.raises(EmptyDocument, match="OCR"):
        parse_document("scan.pdf", doc.tobytes())


def test_cp1252_fallback():
    # "café" encoded in Windows-1252 (0xE9 is invalid UTF-8)
    pages = parse_document("note.txt", b"caf\xe9 gouvernance")
    assert pages == ["café gouvernance"]


def test_invalid_encoding_rejected():
    from app.services.parsing import InvalidEncoding

    # 0x81/0x8D/0x90 are undefined in both UTF-8 (continuation bytes) and cp1252
    with pytest.raises(InvalidEncoding):
        parse_document("binaire.txt", b"\x81\x8d\x90")
