"""Document parsing: bytes -> list of page texts.

PDF via PyMuPDF (one entry per page), DOCX via python-docx (single entry,
paragraphs and tables in true document order), plain text / markdown as-is.
Page numbers are 1-based and serve as the provenance anchor for chunks and
citations downstream — which is why document order must be preserved exactly
and why documents with no extractable text are rejected instead of stored.
"""

import io
import zipfile

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph


class UnsupportedFileType(Exception):
    pass


class EmptyDocument(Exception):
    pass


class InvalidEncoding(Exception):
    pass


class DocumentTooLarge(Exception):
    """A DOCX whose declared uncompressed size exceeds the expansion cap
    (zip-bomb guard): the on-disk file passed the upload limit but would
    inflate far beyond it in memory."""


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}

# Cap on total DECOMPRESSED DOCX content. A .docx is a zip; a small upload can
# declare gigabytes of expanded parts. Checked against the central directory's
# declared sizes — a first-line guard against the common inflate-bomb, not a
# defence against a crafted archive that understates its sizes.
MAX_DOCX_UNCOMPRESSED = 100 * 1024 * 1024  # 100 MB

# Bump whenever extraction logic changes: stored on each Document so M2 can
# detect stale parses and reindex idempotently.
PARSER_VERSION = "2"


def parse_document(filename: str, data: bytes) -> list[str]:
    """Return the extracted text of each page.

    Raises UnsupportedFileType, EmptyDocument or InvalidEncoding.
    """
    name = filename.lower()
    if name.endswith(".pdf"):
        pages = _parse_pdf(data)
    elif name.endswith(".docx"):
        pages = _parse_docx(data)
    elif name.endswith((".txt", ".md")):
        pages = [_decode_text(data)]
    else:
        raise UnsupportedFileType(f"Type de fichier non supporté : {filename}")

    if not any(page.strip() for page in pages):
        hint = " (document scanné ? L'OCR n'est pas pris en charge.)" if name.endswith(".pdf") else ""
        raise EmptyDocument(f"Aucun texte exploitable extrait de {filename}.{hint}")
    return pages


def _decode_text(data: bytes) -> str:
    """UTF-8 first, cp1252 fallback; never silently replaces bytes."""
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return data.decode("cp1252")
        except UnicodeDecodeError as exc:
            raise InvalidEncoding("Encodage du texte non reconnu (UTF-8 ou Windows-1252 attendu).") from exc


def _parse_pdf(data: bytes) -> list[str]:
    with fitz.open(stream=data, filetype="pdf") as doc:
        return [page.get_text("text") for page in doc]


def _parse_docx(data: bytes) -> list[str]:
    """Extract paragraphs and tables in true document order (body XML order)."""
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        expanded = sum(info.file_size for info in zf.infolist())
    if expanded > MAX_DOCX_UNCOMPRESSED:
        raise DocumentTooLarge(
            f"Document DOCX trop volumineux une fois décompressé "
            f"({expanded // (1024 * 1024)} Mo, limite {MAX_DOCX_UNCOMPRESSED // (1024 * 1024)} Mo)."
        )
    doc = DocxDocument(io.BytesIO(data))
    parts: list[str] = []
    for child in doc.element.body.iterchildren():
        if child.tag.endswith("}p"):
            parts.append(Paragraph(child, doc).text)
        elif child.tag.endswith("}tbl"):
            for row in Table(child, doc).rows:
                parts.append("\t".join(cell.text for cell in row.cells))
    return ["\n".join(parts)]
